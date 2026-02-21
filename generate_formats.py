# SCRIPT DE CONSOLIDACIÓN MULTIFORMATO - HMO AUDITOR ELITE
import os
from docx import Document
import openpyxl
from openpyxl.styles import Font

def export_to_word(title, content, filename):
    doc = Document()
    doc.add_heading(title, 0)
    for line in content.split('\n'):
        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        else: doc.add_paragraph(line)
    path = os.path.join(r"C:\Users\Heymolqs\.gemini\antigravity\brain\a9a36189-27f6-4795-8d2a-4af8e023fe84", filename)
    doc.save(path)
    return path

def export_to_excel(title, data, filename):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]
    ws.append([title])
    ws['A1'].font = Font(bold=True, size=14)
    for row in data:
        ws.append(row)
    path = os.path.join(r"C:\Users\Heymolqs\.gemini\antigravity\brain\a9a36189-27f6-4795-8d2a-4af8e023fe84", filename)
    wb.save(path)
    return path

# Plan
plan_txt = """# PLAN DE IMPLEMENTACIÓN MAESTRO - ELITE V1.5.3\n## 🛡️ Arquitectura\n- Motor de Reconocimiento Local.\n- Evaluador Experto.\n## 📊 Fases\n1. Fase A: Identidad.\n2. Fase B: Dimensionamiento.\n3. Fase C: Cuerpo Normativo."""
export_to_word("Plan de Implementación Elite", plan_txt, "Implementacion_Elite_V153.docx")

# Matriz
matriz_data = [["NORMA", "ESTADO"], ["ISO 9001", "ANCLADO"], ["ISO 19011", "ANCLADO"]]
export_to_excel("Matriz Normativa Elite", matriz_data, "Matriz_Normativa_HMO.xlsx")

# Manual
manual_txt = """# MANUAL DE OPERACIÓN\n1. Ingesta\n2. Dashboard\n3. Revisión\n4. Emisión"""
export_to_word("Manual de Operación Elite", manual_txt, "Manual_Usuario_HMO_Elite.docx")

# Innovatech
inv_txt = """# EXPEDIENTE INNOVATECH\nMISIÓN: Rigor digital.\nVISIÓN: Referente 2030."""
export_to_word("Expediente Innovatech", inv_txt, "Innovatech_Mision_Vision.docx")

print("Exportación Multiformato Completada.")
