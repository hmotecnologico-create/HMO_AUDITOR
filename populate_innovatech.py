# SCRIPT DE GENERACIÓN DE EXPEDIENTE RÍGIDO - INNOVATECH SOLUTIONS
import os
from docx import Document
import openpyxl
from openpyxl.styles import Font

base_path = r"d:\HMO\SENA\Auditor_Formatos\Innovatech_Solutions"

def export_to_word(title, content, folder, filename):
    doc = Document()
    doc.add_heading(title, 0)
    for line in content.split('\n'):
        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        else: doc.add_paragraph(line)
    path = os.path.join(base_path, folder, filename)
    doc.save(path)
    return path

def export_to_excel(title, data, folder, filename):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]
    ws.append([title])
    ws['A1'].font = Font(bold=True, size=14)
    for row in data:
        ws.append(row)
    path = os.path.join(base_path, folder, filename)
    wb.save(path)
    return path

# --- 01. DIRECCIÓN Y ESTRATEGIA ---
estrategia_txt = """# PLATAFORMA ESTRATÉGICA 2026
## MISIÓN
Impulsar la eficiencia operativa mediante tecnología de punta.
## VISIÓN
Liderar el sector servicios tecnológicos en el país para 2030.
## VALORES
Excelencia, Transparencia e Innovación Constante.
"""
export_to_word("Plataforma Estratégica", estrategia_txt, "01_Direccion_y_Estrategia", "Mision_Vision_Valores.docx")
export_to_word("Organigrama Nominal", "# ORGANIGRAMA\nGeneral -> Operaciones -> Calidad -> RRHH", "01_Direccion_y_Estrategia", "Organigrama_Corporativo.docx")

# --- 02. GESTIÓN DE CALIDAD ---
calidad_txt = """# MANUAL DE CALIDAD ISO 9001:2015
## ALCANCE
Prestación de servicios de consultoría tecnológica y auditoría.
## POLÍTICA DE CALIDAD
Satisfacción total del cliente mediante mejora continua.
"""
export_to_word("Manual de Calidad", calidad_txt, "02_Gestion_de_Calidad", "Manual_de_Calidad_V1.docx")
matriz_riesgos = [
    ["RIESGO", "IMPACTO", "PROBABILIDAD", "MITIGACIÓN"],
    ["Pérdida de Datos", "Alto", "Baja", "Backup Diario"],
    ["Falla de Red", "Medio", "Media", "Dual-ISP"],
    ["Error Humano", "Alto", "Media", "Capacitación Continua"]
]
export_to_excel("Matriz de Riesgos", matriz_riesgos, "02_Gestion_de_Calidad", "Matriz_Riesgos_HMO.xlsx")

# --- 03. OPERACIONES ---
ops_txt = """# PROCEDIMIENTO OPERATIVO ESTÁNDAR (SOP)
## OBJETIVO
Estandarizar la ejecución de servicios de soporte.
## PASOS
1. Recepción de Ticket.
2. Análisis de Motor Experto.
3. Validación Humana.
4. Cierre y Trazabilidad.
"""
export_to_word("SOP de Servicios", ops_txt, "03_Operaciones", "Procedimiento_Operativo_Estandar.docx")

# --- 04. RECURSOS HUMANOS ---
rrhh_txt = """# MANUAL DE FUNCIONES
## CARGO: AUDITOR LÍDER
- Responsable de la integridad del expediente.
- Validador final de plantillas base.
- Supervisor del Motor de Reconocimiento.
"""
export_to_word("Manual de Funciones", rrhh_txt, "04_Recursos_Humanos", "Manual_Funciones_Competencias.docx")

print("Expediente Rígido de Innovatech Generado Exitosamente por Procesos.")
