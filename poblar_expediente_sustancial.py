# SCRIPT DE RE-POBLAMIENTO DETERMINÍSTICO - INNOVATECH SOLUTIONS SAS V2.0
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill

base_path = r"d:\HMO\SENA\Auditor_Formatos\Innovatech_Solutions"

def generate_substantial_doc(title, folder, filename, sections):
    doc = Document()
    
    # Header Estilizado
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 2, width=Inches(6))
    htable.cell(0, 0).text = "HMO AUDITOR - ELITE EDITION"
    htable.cell(0, 1).text = f"REF: {filename.split('.')[0]}"
    htable.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Título Principal
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cuerpo del Documento (Párrafos Sustanciales)
    for subtitle, content in sections.items():
        h = doc.add_heading(subtitle, level=1)
        h.runs[0].font.color.rgb = RGBColor(0, 194, 255)
        
        # Simular párrafos largos y densos
        p = doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Añadir sub-secciones para mayor densidad
        doc.add_paragraph("Detalles Técnicos y Justificación:", style='List Bullet')
        doc.add_paragraph(f"Este componente de {subtitle.lower()} se fundamenta en los pilares de la Industria 4.0, integrando protocolos de seguridad redundantes y una arquitectura de datos optimizada para el Motor de Reconocimiento del sistema HMO.", style='List Bullet')

    path = os.path.join(base_path, folder, filename)
    doc.save(path)
    print(f"Generado: {path} (Tamaño estimado: {len(content)} chars)")
    return path

# --- POBLAR DIRECCIÓN Y ESTRATEGIA ---
mision_data = {
    "DECLARACIÓN DE MISIÓN": "Innovatech Solutions SAS tiene como misión fundamental el diseño, desarrollo y manufactura de ecosistemas tecnológicos de alta precisión, integrando Inteligencia Artificial Nativa y hardware de procesamiento distribuido. Nos enfocamos en proveer a la industria manufacturera y de servicios herramientas de auditoría y control de calidad que eliminen el error humano, optimizando la cadena de valor mediante la digitalización absoluta y la veracidad de los datos en tiempo real. Cada solución entregada es el resultado de un riguroso proceso de ingeniería que garantiza la soberanía tecnológica de nuestros clientes.",
    "FUNDAMENTOS OPERATIVOS": "Operamos bajo un modelo de excelencia técnica donde la innovación no es una opción, sino el motor de nuestra existencia. Nuestra misión se materializa en la reducción sistemática de costos operativos para nuestros aliados mediante la implementación de algoritmos de predicción y sistemas de control in situ que operan 24/7 con integridad certificada.",
    "VISIÓN COMPROMETIDA": "Buscamos ser el socio estratégico indispensable para cualquier organización que aspire a la excelencia normativa, convirtiendo la complejidad de los marcos legales en rutas claras de ejecución digital."
}
generate_substantial_doc("MISIÓN CORPORATIVA V2.0", "01_Direccion_y_Estrategia", "01_Mision_Corporativa.docx", mision_data)

vision_data = {
    "VISIÓN ESTRATÉGICA 2030": "Para el año 2030, Innovatech Solutions SAS será reconocida como la empresa líder en Latinoamérica en el desarrollo de Gemelos Digitales y sistemas de Auditoría Inteligente. Aspirasmos a consolidar una red de infraestructuras tecnológicas soberanas que permitan a los gobiernos y empresas privadas operar con niveles de transparencia y eficiencia del 99.9%. Nuestra visión incluye la expansión a 15 países, liderando la transición hacia la Industria 5.0 donde el talento humano y la tecnología disruptiva conviven en armonía perfecta para el beneficio social.",
    "METAS DE ALTO IMPACTO": "1. Internacionalización de la suite HMO Auditor en mercados europeos y asiáticos. \n2. Desarrollo del primer microprocesador de arquitectura abierta para seguridad industrial en la región. \n3. Carbono Neutralidad absoluta mediante el uso de hardware de bajo consumo y centros de datos biosostenibles.",
    "POSICIONAMIENTO GLOBAL": "Innovatech no solo vende productos, construye el futuro de la confianza digital a través de la superioridad técnica y el rigor ético innegociable."
}
generate_substantial_doc("VISIÓN CORPORATIVA V2.0", "01_Direccion_y_Estrategia", "02_Vision_Corporativa.docx", vision_data)

# --- VALORES ---
valores_data = {
    "NÚCLEO AXIOLÓGICO": "1. RIGOR ANALÍTICO: No aceptamos aproximaciones; nuestra ingeniería es exacta. \n2. TRANSPARENCIA RADICAL: Cada proceso es auditable y trazable mediante SHA-256. \n3. INNOVACIÓN SIN FRONTERAS: Desafiamos el 'statu quo' tecnológico diariamente. \n4. COMPROMISO ÉTICO: La privacidad y soberanía de los datos son nuestros pilares.",
    "CULTURA DE ÉLITE": "Fomentamos un entorno de aprendizaje continuo donde el fracaso es una lección de ingeniería y el éxito es un estándar que se supera cada día."
}
generate_substantial_doc("VALORES NUCLEARES V2.0", "01_Direccion_y_Estrategia", "03_Values_Elite.docx", valores_data)

# --- CALIDAD ---
calidad_data = {
    "SISTEMA DE GESTIÓN DE CALIDAD (SGC)": "El SGC de Innovatech Solutions SAS se basa en la norma ISO 9001:2015, extendida con protocolos de seguridad ISO 27001. Nuestra metodología de desarrollo sigue el estándar CMMI Nivel 5, asegurando que cada fase del diseño, desde la ingeniería conceptual hasta el despliegue final, sea monitorizada por nuestro Evaluador Experto. La calidad no es un departamento, es el resultado de una ejecución perfecta en cada punto del mapa de procesos.",
    "POLÍTICA DE CALIDAD INTEGRAL": "Nos comprometemos a entregar activos tecnológicos que no solo cumplan con los requisitos del cliente, sino que superen sus expectativas de madurez digital. Nuestra política se centra en la Prevención sobre la Corrección, utilizando el Motor de Reconocimiento del HMO para identificar desviaciones normativas en milisegundos."
}
generate_substantial_doc("MANUAL DE CALIDAD ELITE", "02_Gestion_de_Calidad", "Manual_de_Calidad_V1.docx", calidad_data)

print("EXPEDIENTE RE-POBLADO CON DENSIDAD DE ÉLITE V2.0")
