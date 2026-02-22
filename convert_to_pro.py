import os
from docx import Document

base_path = r"d:\HMO\SENA\Auditor_Formatos\Auditorias_HMO\Innovatech_Solutions"

files_to_convert = [
    r"01_Direccion_y_Estrategia\Acta_Reunion_Directiva_Inicio.md",
    r"01_Direccion_y_Estrategia\Mision_y_Vision.md",
    r"02_Gestion_de_Calidad\Politica_de_Calidad.md",
    r"03_Operaciones\Mapa_de_Procesos.md"
]

def convert_md_to_docx(rel_path):
    abs_path = os.path.join(base_path, rel_path)
    if not os.path.exists(abs_path):
        print(f"File not found: {abs_path}")
        return

    with open(abs_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    doc.add_heading(rel_path.split('\\')[-1].replace('.md', '').replace('_', ' '), 0)
    
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.strip():
            doc.add_paragraph(line)

    output_path = abs_path.replace('.md', '.docx')
    doc.save(output_path)
    print(f"Converted: {output_path}")
    # os.remove(abs_path) # Comentado por seguridad hasta verificar

if __name__ == "__main__":
    for f in files_to_convert:
        convert_md_to_docx(f)
