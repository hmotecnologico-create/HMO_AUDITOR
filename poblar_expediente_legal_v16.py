# SCRIPT DE GENERACIÓN DE EXPEDIENTE LEGAL V1.6.0 - INNOVATECH SOLUTIONS SAS
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

base_path = r"d:\HMO\SENA\Auditor_Formatos\Innovatech_Solutions"

def generate_legal_doc(title, folder, filename, sections):
    doc = Document()
    
    # Header Estilizado (Rigor Legal)
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 2, width=Inches(6))
    htable.cell(0, 0).text = "REPÚBLICA DE COLOMBIA - SISTEMA DE REGISTRO"
    htable.cell(0, 1).text = f"DOCUMENTO: {filename.split('.')[0]}"
    htable.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Título Principal
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for subtitle, content in sections.items():
        h = doc.add_heading(subtitle, level=1)
        h.runs[0].font.color.rgb = RGBColor(0, 32, 96)
        p = doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.style.font.size = Pt(10)

    path = os.path.join(base_path, folder, filename)
    doc.save(path)
    print(f"Generado Legal: {path}")
    return path

# --- 04. CÁMARA DE COMERCIO SIMULADA ---
camara_data = {
    "DATOS DE LA ENTIDAD": "RAZÓN SOCIAL: INNOVATECH SOLUTIONS SAS\nNIT: 901.455.789-2\nDOMICILIO: Calle 100 #7-33, Edificio Synergy Tower, Piso 15, Bogotá D.C., Colombia.\nTELÉFONO: (+57) 601 789 4455\nWEB: www.innovatechsolutions.com.co",
    "OBJETO SOCIAL PRINCIPAL": "La sociedad tiene como objeto principal el diseño, desarrollo, implementación, comercialización y mantenimiento de soluciones integrales de software avanzado, sistemas de inteligencia artificial, y la fabricación, ensamble e importación de hardware especializado para la automatización industrial y auditoría digital. Así mismo, podrá prestar servicios de consultoría técnica de alto nivel y capacitación en transformación digital.",
    "REPRESENTACIÓN LEGAL": "Representante Legal Titular: Ing. Carlos Martínez Serna (C.C. 1.018.445.678).\nFacultades: Plenas para la gestión administrativa, técnica y financiera hasta la cuantía de 10,000 SMLV.",
    "CERTIFICACIÓN DE EXISTENCIA": "La sociedad fue constituida mediante acta No. 001 de fecha 12 de Enero de 2022, debidamente inscrita en el registro mercantil del domicilio principal."
}
generate_legal_doc("CERTIFICADO DE EXISTENCIA Y REPRESENTACIÓN LEGAL", "01_Direccion_y_Estrategia", "04_Camara_Comercio_Simulada.docx", camara_data)

# --- 05. RUT SIMULADO ---
rut_data = {
    "IDENTIFICACIÓN TRIBUTARIA": "NIT: 901.455.789-2\nESTADO: Activo\nACTIVIDAD ECONÓMICA PRINCIPAL: 6201 (Actividades de desarrollo de sistemas informáticos).\nACTIVIDAD SECUNDARIA: 2610 (Fabricación de componentes y tableros electrónicos).",
    "RESPONSABILIDADES FISCALES": "05 - Impuesto sobre la Renta y Complementarios.\n07 - Retención en la fuente a título de renta.\n14 - Informante de exógena.\n42 - Obligado a llevar contabilidad.",
    "DIRECCIÓN NOTIFICACIONES": "Calle 100 #7-33, Bogotá D.C. - Correo: tributario@innovatechsolutions.com.co"
}
generate_legal_doc("REGISTRO ÚNICO TRIBUTARIO (SIMULADO)", "01_Direccion_y_Estrategia", "05_RUT_Simulado.docx", rut_data)

# --- 06. MATRIZ DE RESPONSABLES ---
responsables_data = {
    "ESTRUCTURA DE GOBERNANZA": "Para efectos de la auditoría HMO Elite, se definen los siguientes responsables de área con autoridad sobre el expediente documental:",
    "DIRECCIÓN GENERAL Y ESTRATEGIA": "Responsable: Ing. Carlos Martínez Serna (CEO)\nAlcance: Aprobación de presupuestos y visión estratégica.",
    "GESTIÓN DE CALIDAD Y NORMATIVIDAD": "Responsable: Dra. Laura Giraldo Espitia (Directora de Calidad)\nAlcance: Integridad del Manual de Calidad y cumplimiento ISO 9001/27001.",
    "OPERACIONES Y MANUFACTURA": "Responsable: Ing. Carlos Martínez (Acting Ops Manager)\nAlcance: Supervisión de líneas de producción y desarrollo de software.",
    "TALENTO HUMANO Y CULTURA": "Responsable: Lic. Roberto Gómez (Gerente de Talento)\nAlcance: Manual de funciones y competencias del personal auditado."
}
generate_legal_doc("MATRIZ DE RESPONSABILIDADES Y LIDERAZGO", "01_Direccion_y_Estrategia", "06_Matriz_Responsables_Area.docx", responsables_data)

print("\n--- EXPEDIENTE LEGAL V1.6.0 CONSOLIDADO ---")
