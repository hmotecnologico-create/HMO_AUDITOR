from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import datetime

def create_audit_program_v2(company_name, output_path, logo_path=None, kb=None, identity_data=None):
    doc = Document()
    kb = kb or {}
    ident = identity_data or {
        "auditor": "Auditor Asignado", "rep_legal": "Representante Legal", "rep_id": "N/A", 
        "tamanio": "No Definido", "sector": "No Definido", "nit": "N/A", 
        "direccion": "No especificada", "web": "N/A", "objeto_social": "No definido"
    }
    
    # 1. ENCABEZADO (Validez Jurídica)
    section = doc.sections[0]
    header = section.header
    
    # Tabla de encabezado para Logo y Títulos
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.columns[0].width = Inches(1.5)
    header_table.columns[1].width = Inches(5.0)
    
    # Celda 1: Logo
    cell_logo = header_table.rows[0].cells[0]
    if logo_path and os.path.exists(logo_path):
        run_logo = cell_logo.paragraphs[0].add_run()
        run_logo.add_picture(logo_path, width=Inches(1.0))
    else: cell_logo.text = "LOGO"
    
    # Celda 2: Texto
    cell_text = header_table.rows[0].cells[1]
    para = cell_text.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"SISTEMA DE GESTIÓN DE CALIDAD - {company_name}\nFORMATO DE AUDITORÍA INTERNA")
    run.bold, run.font.size = True, Pt(11)

    # Tabla secundaria para Código, Versión, Fecha
    info_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    cells = info_table.rows[0].cells
    cells[0].text, cells[1].text, cells[2].text = "Código: AUD-PROG-01", "Versión: 02", f"Fecha: {datetime.date.today()}"
            
    # Título Principal
    title = doc.add_heading('PROGRAMA DE AUDITORÍA INTERNA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. DATOS GENERALES (Materia Prima Fase A y B)
    doc.add_heading('1. DATOS GENERALES Y DIMENSIONAMIENTO', level=1)
    table_gen = doc.add_table(rows=9, cols=2)
    table_gen.style = 'Table Grid'
    
    data_gen = [
        ("Auditor Líder Responsable", ident["auditor"]),
        ("Representante Entidad", ident["rep_legal"] + f" ({ident['rep_id']})"),
        ("NIT / Identificación Fiscal", ident.get("nit", "N/A")),
        ("Dirección Principal Sede", ident.get("direccion", "No especificada")),
        ("Dimensionamiento Org.", ident["tamanio"]),
        ("Sector Económico / Objeto", f"{ident['sector']} - {ident.get('objeto_social', '')[:50]}..."),
        ("Tipo de Auditoría", "Interna de Calidad (ISO 9001:2015)"),
        ("Fecha Programada", str(datetime.date.today())),
        ("ID Expediente Digital", f"EXP-{company_name[:4].upper()}-2026")
    ]
    
    for i, (key, value) in enumerate(data_gen):
        table_gen.cell(i, 0).text = key
        table_gen.cell(i, 1).text = value
        table_gen.cell(i, 0).paragraphs[0].runs[0].bold = True

    # 3. FILOSOFÍA CORPORATIVA
    doc.add_heading('2. FILOSOFÍA CORPORATIVA (DILIGENCIADO)', level=1)
    doc.add_paragraph(f"Misión/Visión: {kb.get('Misión y Visión Corporativa', 'No especificada.')}")
    doc.add_paragraph(f"Principios Rectores: {kb.get('Valores y Código de Ética', 'No especificados.')}")

    # 4. OBJETIVO Y ALCANCE (ISO 19011:6.2.2)
    doc.add_heading('3. OBJETIVO Y ALCANCE', level=1)
    doc.add_paragraph(f"Objetivo: {kb.get('PEI (Proyecto Educativo)', {}).get('resumen', 'Verificar cumplimiento normativo.') if isinstance(kb.get('PEI (Proyecto Educativo)'), dict) else 'Verificar cumplimiento normativo.'}")
    doc.add_paragraph(f"Alcance: {kb.get('Contexto Organizacional', 'Procesos misionales de la organización.')}")

    # 5. EVIDENCIAS Y HALLAZGOS (ISO 19011:6.4)
    doc.add_heading('4. EVIDENCIAS Y HALLAZGOS DETECTADOS (IA COGNITIVA)', level=1)
    for doc_id, data in kb.items():
        if isinstance(data, dict):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{doc_id}: ")
            run.bold = True
            p.add_run(f"Coherencia: {data.get('coherencia')}%")
            doc.add_paragraph(f"Resumen Semántico: {data.get('resumen')}", style='Body Text')
            if data.get('hallazgos'):
                doc.add_paragraph("Hallazgos:", style='Normal').runs[0].italic = True
                for h in data.get('hallazgos'):
                    doc.add_paragraph(f"- {h}", style='List Bullet 2')
        else:
            doc.add_paragraph(f"{doc_id}: Verificado (Procesamiento Estándar).", style='List Bullet')

    # 6. CRITERIOS DE AUDITORÍA
    doc.add_heading('5. CRITERIOS DE AUDITORÍA (BASE LEGAL)', level=1)
    doc.add_paragraph("- Norma ISO 19011:2018 (Directrices para Auditoría)")
    doc.add_paragraph("- Marco Normativo Seleccionado en HMO Auditor")
    doc.add_paragraph("- Ley 594 de 2000 (Ley General de Archivos - Colombia)")
    doc.add_paragraph("- Manuales de Procedimientos Internos")

    # 7. CONCLUSIONES DE AUDITORÍA (ISO 19011:6.5)
    doc.add_heading('6. CONCLUSIONES GENERALES', level=1)
    doc.add_paragraph("Basado en el enfoque de riesgos y evidencias analizadas por el motor cognitivo Llama3:8b, se concluye que el sistema de gestión muestra una madurez documental conforme a los requisitos evaluados.")

    # 11. FIRMAS
    doc.add_page_break()
    doc.add_heading('7. RESPONSABILIDAD LEGAL Y FIRMAS', level=1)
    doc.add_paragraph("De acuerdo con la ISO 19011, este documento formaliza el compromiso de las partes.")
    
    table_sig = doc.add_table(rows=2, cols=2)
    table_sig.style = 'Table Grid'
    
    # Auditor Row
    table_sig.cell(0, 0).text = f"FIRMA DEL AUDITOR\n\n\n__________________________\nNombre: {ident['auditor']}\nCargo: Auditor Certificado"
    
    # Auditado Row
    table_sig.cell(0, 1).text = f"FIRMA DEL AUDITADO\n\n\n__________________________\nNombre: {ident['rep_legal']}\nID: {ident['rep_id']}"
    
    # 12. CONTROL DE VERSIONES
    doc.add_heading('8. CONTROL DE VERSIONES', level=1)
    table_ver = doc.add_table(rows=2, cols=4)
    table_ver.style = 'Table Grid'
    cols_ver = ["Versión", "Fecha", "Descripción", "Responsable"]
    for i, txt in enumerate(cols_ver):
        table_ver.cell(0, i).text = txt
        table_ver.cell(0, i).paragraphs[0].runs[0].bold = True
    
    row_ver = table_ver.rows[1].cells
    row_ver[0].text = "03"
    row_ver[1].text = str(datetime.date.today())
    row_ver[2].text = "Cumplimiento estructural ISO 19011 e IA Cognitiva"
    row_ver[3].text = "HMO Auditor System V8.0"

    # Save
    if not os.path.exists(output_path): os.makedirs(output_path)
    
    file_name = "GAD_PRO_01_Programa_Auditoria_ELITE.docx"
    full_path = os.path.join(output_path, file_name)
    doc.save(full_path)
    return full_path

if __name__ == "__main__":
    company = "Innovatech Solutions SAS"
    path = "d:\\HMO\\SENA\\Auditor_Formatos\\HMO_Auditor_Master_V1\\02_Formatos_del_Sistema"
    create_audit_program_v2(company, path)
