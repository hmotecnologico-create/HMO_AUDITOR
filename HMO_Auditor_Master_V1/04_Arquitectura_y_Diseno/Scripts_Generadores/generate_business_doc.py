from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_business_proposal():
    doc = Document()
    
    # Estilo de Título
    title = doc.add_heading('HMO AUDITOR ELITE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Propuesta de Negocio y Estrategia de Crecimiento', 1)
    
    p = doc.add_paragraph()
    p.add_run('Versión 1.3 - Edición Profesional').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading('1. Propuesta de Valor Única (UVP)', 2)
    doc.add_paragraph(
        "HMO Auditor revoluciona el sector de la auditoría mediante la integración de Inteligencia Artificial (IA) "
        "operando 100% en local. Resolvemos el conflicto entre eficiencia y confidencialidad: "
        "los datos sensibles nunca salen de la infraestructura del cliente."
    )

    doc.add_heading('2. Modelo de Negocio (Monetización)', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Modelo'
    hdr_cells[1].text = 'Descripción'
    hdr_cells[2].text = 'Potencial'
    
    data = [
        ('Licencia Anual', 'Venta por puesto de trabajo para firmas de auditoría.', 'Ingreso Recurrente'),
        ('Consultoría RAG', 'Configuración de bases de conocimiento personalizadas.', 'Pago por Proyecto'),
        ('Soporte Elite', 'Actualizaciones normativas periódicas (ISO/MEN).', 'Suscripción Mensual')
    ]
    
    for mod, desc, pot in data:
        row_cells = table.add_row().cells
        row_cells[0].text = mod
        row_cells[1].text = desc
        row_cells[2].text = pot

    doc.add_heading('3. Estrategia de Marketing e Innovación', 2)
    doc.add_paragraph(
        "Nuestra estrategia se centra en el 'Océano Azul' de la privacidad. "
        "Atacamos el sector académico (Registro Calificado) y el corporativo (ISO) "
        "posicionándonos como la única alternativa que ofrece IA sin riesgo de fuga de datos."
    )
    
    doc.add_paragraph("- Focus: Instituciones Educativas y Firmas de Auditoría ISO.")
    doc.add_paragraph("- Canal: LinkedIn Thought Leadership y Alianzas con Certificadoras.")

    doc.add_heading('4. Opinión Estratégica y Futuro', 2)
    doc.add_paragraph(
        "El ecosistema HMO Auditor está listo para escalar. La V2.0 incluirá Visión Artificial "
        "y certificación de integridad vía Blockchain, posicionando el producto al 1000% de su capacidad comercial."
    )

    output_path = r'd:\HMO\SENA\Auditor_Formatos\HMO_Auditor_Master_V1\06_Documentacion_Elite\Propuesta_de_Negocio_HMO_Elite.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Propuesta de Negocio generada en: {output_path}")

if __name__ == "__main__":
    create_business_proposal()
