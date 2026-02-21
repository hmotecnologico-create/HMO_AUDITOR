from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_normative_sources_docx():
    doc = Document()
    
    title = doc.add_heading('MATRIZ DE FUENTES Y VERACIDAD NORMATIVA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('HMO Auditor Elite - Root of Trust', 1)
    
    doc.add_heading('1. Fuentes Oficiales de Conocimiento', 2)
    doc.add_paragraph(
        "La base de conocimiento del sistema se sustenta en fuentes públicas y estándares internacionales "
        "para garantizar la validez del análisis de auditoría."
    )
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sistema'
    hdr_cells[1].text = 'Referencia Principal'
    hdr_cells[2].text = 'Origen'
    
    data = [
        ("SGC", "ISO 9001:2015", "Estándar Internacional"),
        ("SGSI", "ISO 27001:2022", "Estándar Internacional"),
        ("Académico", "Ley 115 / Dec. 1330", "MEN Colombia"),
        ("Auditoría", "ISO 19011:2018", "Directrices Internacionales")
    ]
    
    for sys, ref, src in data:
        row_cells = table.add_row().cells
        row_cells[0].text = sys
        row_cells[1].text = ref
        row_cells[2].text = src

    doc.add_heading('2. Protocolo de No Alucinación (RAG)', 2)
    doc.add_paragraph(
        "El motor RAG prioriza los fragmentos de los documentos cargados sobre el conocimiento general del LLM, "
        "asegurando que cada sugerencia tenga una cita normativa exacta."
    )

    output_path = r'd:\HMO\SENA\Auditor_Formatos\HMO_Auditor_Master_V1\06_Documentacion_Elite\Matriz_Fuentes_Normativas_Elite.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Normative Sources docx generated at: {output_path}")

if __name__ == "__main__":
    create_normative_sources_docx()
