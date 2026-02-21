from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_audit_program(company_name, output_path):
    doc = Document()
    
    # --- Header ---
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"HMO AUDITOR - SISTEMA DE GESTIÓN DE CALIDAD\nCÓDIGO: GAD-PROG-01 | VERSIÓN: 01"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- Title ---
    title = doc.add_heading('PROGRAMA DE AUDITORÍA INTERNA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- Company Info Table ---
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = "Empresa:"
    table.cell(0, 1).text = company_name
    table.cell(1, 0).text = "Ciclo de Auditoría:"
    table.cell(1, 1).text = "Anual 2026"
    table.cell(2, 0).text = "Auditor Líder:"
    table.cell(2, 1).text = "Ing. Roberto Steiner (Asig. AI)"
    
    doc.add_paragraph("\n")
    
    # --- Objectives and Scope (AI Diligenced) ---
    doc.add_heading('1. OBJETIVOS Y ALCANCE', level=1)
    p = doc.add_paragraph()
    run = p.add_run("Objetivo: ")
    run.bold = True
    p.add_run("Evaluar el cumplimiento de los procesos de Innovatech Solutions SAS frente a la norma ISO 9001:2015, asegurando que la innovación tecnológica se mantenga bajo estándares de calidad controlados.")
    
    p = doc.add_paragraph()
    run = p.add_run("Alcance: ")
    run.bold = True
    p.add_run("Incluye todos los procesos del Mapa de Procesos V.01, específicamente el área de Desarrollo de Software y R&D.")
    
    # --- Audit Schedule ---
    doc.add_heading('2. CRONOGRAMA DE ACTIVIDADES', level=1)
    sched_table = doc.add_table(rows=4, cols=3)
    sched_table.style = 'Table Grid'
    
    hdr_cells = sched_table.rows[0].cells
    hdr_cells[0].text = 'Proceso'
    hdr_cells[1].text = 'Fecha Sugerida AI'
    hdr_cells[2].text = 'Estado'
    
    processes = [
        ["Dirección Estratégica", "15/03/2026", "Programado"],
        ["Gestión de Calidad", "20/03/2026", "Programado"],
        ["Desarrollo de Software", "05/04/2026", "Programado"]
    ]
    
    for i, row in enumerate(processes):
        row_cells = sched_table.rows[i+1].cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        
    doc.add_paragraph("\n")
    doc.add_paragraph("Este documento ha sido generado automáticamente por HMO Auditor. La estructura está protegida contra modificaciones no autorizadas conforme a la norma ISO 27001.")
    
    # Final cleanup and save
    full_output = os.path.join(output_path, "GAD_PROG_01_Programa_Auditoria_Innovatech.docx")
    doc.save(full_output)
    print(f"Documento generado en: {full_output}")

if __name__ == "__main__":
    output_dir = r"d:\HMO\SENA\Auditor_Formatos\Formatos_Profesionales_HMO"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    create_audit_program("Innovatech Solutions SAS", output_dir)
