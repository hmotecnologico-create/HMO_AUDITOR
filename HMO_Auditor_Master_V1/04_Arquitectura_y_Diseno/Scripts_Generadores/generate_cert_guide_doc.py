from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_certification_guide_docx():
    doc = Document()
    
    title = doc.add_heading('RUTA A LA CERTIFICACIÓN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Guía Paso a Paso para No Expertos', 1)
    
    doc.add_heading('Fase 1: Cimentación (Mes 1)', 2)
    doc.add_paragraph("Objetivo: Establecer el compromiso de la empresa y los pilares estratégicos.")
    doc.add_paragraph("- Reunión Directiva y Acta de Compromiso.")
    doc.add_paragraph("- Definición de Misión, Visión y Política de Calidad.")
    
    doc.add_heading('Fase 2: Documentación e Ingesta (Mes 2-3)', 2)
    doc.add_paragraph("Objetivo: Alimentar al motor RAG con el contexto organizacional.")
    doc.add_paragraph("- Carga masiva de manuales y procedimientos.")
    doc.add_paragraph("- Validación Humana (HITL) del texto extraído.")
    
    doc.add_heading('Fase 3: Ejecución de Auditorías (Mes 4-5)', 2)
    doc.add_paragraph("Objetivo: Generar y completar las Listas de Verificación asistidas por IA.")
    doc.add_paragraph("- Generación automatizada del Programa de Auditoría.")
    doc.add_paragraph("- Cierre de hallazgos con lenguaje técnico profesional.")
    
    doc.add_heading('Fase 4: Certificación Final (Mes 6)', 2)
    doc.add_paragraph("Objetivo: Presentación del expediente blindado ante entes externos.")

    output_path = r'd:\HMO\SENA\Auditor_Formatos\HMO_Auditor_Master_V1\06_Documentacion_Elite\Guia_de_Certificacion_HMO.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Certification Guide docx generated at: {output_path}")

if __name__ == "__main__":
    create_certification_guide_docx()
