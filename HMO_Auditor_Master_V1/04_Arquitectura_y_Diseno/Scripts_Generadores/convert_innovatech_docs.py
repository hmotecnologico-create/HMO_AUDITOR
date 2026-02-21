from docx import Document
from docx.shared import Pt
import os

def convert_md_to_docx(md_path, docx_path, title):
    doc = Document()
    doc.add_heading(title, 0)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('**'):
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"Convertido: {docx_path}")

if __name__ == "__main__":
    base_dir = r"d:\HMO\SENA\Auditor_Formatos\Innovatech_Solutions"
    output_dir = r"d:\HMO\SENA\Auditor_Formatos\Formatos_Profesionales_HMO\Innovatech_Reales"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    files_to_convert = [
        (os.path.join(base_dir, "01_Direccion_y_Estrategia", "Mision_y_Vision.md"), "Mision_y_Vision.docx", "Misión y Visión Corporativa"),
        (os.path.join(base_dir, "01_Direccion_y_Estrategia", "Acta_Reunion_Directiva_Inicio.md"), "Acta_Inicio_ISO9001.docx", "Acta de Inicio - Certificación"),
        (os.path.join(base_dir, "02_Gestion_de_Calidad", "Politica_de_Calidad.md"), "Politica_de_Calidad.docx", "Política de Calidad V.01"),
        (os.path.join(base_dir, "03_Operaciones", "Mapa_de_Procesos.md"), "Mapa_de_Procesos.docx", "Mapa de Procesos Corporativo")
    ]
    
    for md, docx, title in files_to_convert:
        convert_md_to_docx(md, os.path.join(output_dir, docx), title)
