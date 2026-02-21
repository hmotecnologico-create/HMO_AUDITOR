from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_implementation_plan_docx():
    doc = Document()
    
    # Titulo principal
    title = doc.add_heading('PLAN DE IMPLEMENTACIÓN TÉCNICA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('HMO Auditor V1.3 Elite', 1)
    
    doc.add_heading('1. Arquitectura Técnica (Local-First)', 2)
    doc.add_paragraph(
        "El sistema opera bajo una arquitectura de Generación Aumentada por Recuperación (RAG) "
        "ejecutada íntegramente en el hardware local del usuario. Esto garantiza la soberanía de los datos."
    )
    
    doc.add_heading('1.1. Componentes Core', 3)
    doc.add_paragraph("- Motor LLM: Ollama (Llama 3 / Mistral).")
    doc.add_paragraph("- Vector Store: ChromaDB (Local Persistence).")
    doc.add_paragraph("- UI: Streamlit Dashboard Profesional.")

    doc.add_heading('2. Estrategia de Blindaje Legal', 2)
    doc.add_paragraph(
        "Cada documento emitido (Word, Excel, PDF) incorpora un sello de integridad SHA-256. "
        "Las plantillas están protegidas mediante 'Form Fields' para impedir alteraciones estructurales."
    )

    doc.add_heading('3. Metodología de Auditoría', 2)
    doc.add_paragraph(
        "Alineación total con la norma ISO 19011:2018. El sistema automatiza el Programa de Auditoría, "
        "el Checklist y el Reporte de Hallazgos, permitiendo una validación humana (HITL) en cada etapa."
    )

    output_path = r'd:\HMO\SENA\Auditor_Formatos\HMO_Auditor_Master_V1\06_Documentacion_Elite\Plan_de_Implementacion_HMO_Elite.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Implementation Plan docx generated at: {output_path}")

if __name__ == "__main__":
    create_implementation_plan_docx()
