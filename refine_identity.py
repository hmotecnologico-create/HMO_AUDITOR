# SCRIPT DE REFINAMIENTO DE IDENTIDAD - INNOVATECH SOLUTIONS SAS
import os
from docx import Document
from docx.shared import Pt
import openpyxl
from openpyxl.styles import Font

base_path = r"d:\HMO\SENA\Auditor_Formatos\Innovatech_Solutions"

def export_to_word(title, content, folder, filename):
    doc = Document()
    # Estilo de encabezado elegante
    header = doc.add_heading(title, 0)
    header.alignment = 1 # Centrado
    
    for line in content.split('\n'):
        if line.startswith('# '): 
            h = doc.add_heading(line[2:], level=1)
            h.style.font.size = Pt(16)
        elif line.startswith('## '): 
            h = doc.add_heading(line[3:], level=2)
            h.style.font.size = Pt(14)
        else: 
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)
            
    path = os.path.join(base_path, folder, filename)
    doc.save(path)
    return path

# --- 01. DIRECCIÓN Y ESTRATEGIA (ACTUALIZADO: DOCUMENTOS SEPARADOS) ---

mision_txt = """# MISIÓN CORPORATIVA - INNOVATECH SOLUTIONS SAS
Nuestra razón de ser es el diseño, desarrollo y despliegue de ecosistemas de software de alta complejidad y la fabricación de hardware inteligente que redefine la productividad industrial. 

Nos comprometemos a entregar soluciones tecnológicas de vanguardia que permitan a nuestros clientes liderar sus respectivos mercados, garantizando siempre la trazabilidad, seguridad y eficiencia operativa mediante la mejora continua de nuestros procesos bajo estándares internacionales de calidad.
"""
export_to_word("Misión de la Entidad", mision_txt, "01_Direccion_y_Estrategia", "01_Mision_Corporativa.docx")

vision_txt = """# VISIÓN CORPORATIVA - INNOVATECH SOLUTIONS SAS (Rumbo 2030)
Para el año 2030, Innovatech Solutions SAS será reconocida globalmente como la fábrica de software y hardware más disruptiva de América Latina, siendo el referente obligatorio en auditoría tecnológica automatizada y aseguramiento de la veracidad digital. 

Aspiramos a consolidar una infraestructura de ingeniería que no solo responda a las necesidades del mercado, sino que anticipe los desafíos de la industria 4.0, manteniendo un crecimiento sostenible y una cultura de innovación innegociable.
"""
export_to_word("Visión de la Entidad", vision_txt, "01_Direccion_y_Estrategia", "02_Vision_Corporativa.docx")

# Valores (Separado para mayor rigor)
valores_txt = """# VALORES NUCLEARES - ADN INNOVATECH
1. INNOVACIÓN DISRUPTIVA: Cuestionamos lo establecido para crear lo extraordinario.
2. RIGOR TÉCNICO: Nuestra ingeniería es exacta, auditable y de calidad superior.
3. ÉTICA DIGITAL: Garantizamos la soberanía y seguridad de los datos de nuestros aliados.
4. AGILIDAD ESTRATÉGICA: Respondemos con precisión y velocidad a los cambios del entorno global.
"""
export_to_word("Valores Corporativos", valores_txt, "01_Direccion_y_Estrategia", "03_Valores_Nucleares.docx")

# Limpieza de archivo anterior combinado
old_file = os.path.join(base_path, "01_Direccion_y_Estrategia", "Mision_Vision_Valores.docx")
if os.path.exists(old_file):
    os.remove(old_file)

print("Identidad de Innovatech Refinada: Misión, Visión y Valores separados con contenido profesional.")
