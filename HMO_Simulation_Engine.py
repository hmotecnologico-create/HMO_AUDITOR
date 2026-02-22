import os
import json
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from HMO_AI_Engine import HMO_AI_Engine

class HMOSimulationEngine:
    def __init__(self, base_path, ai_engine=None):
        self.base_path = base_path
        self.ai_engine = ai_engine or HMO_AI_Engine()
        self.norms_config = {
            "CALIDAD": ["ISO 9001:2015", "Manual de Calidad", "Mapa de Procesos", "Matriz de Riesgos SGC"],
            "SEGURIDAD": ["ISO 27001:2022", "Politica de Seguridad", "SoA (Declaracion de Aplicabilidad)", "Analisis de Riesgos TI"],
            "AMBIENTAL": ["ISO 14001:2015", "Aspectos e Impactos Ambientales", "Matriz de Requisitos Legales", "Plan de Emergencias"],
            "ACADEMICO": ["Ley 115 / Dec 1330", "PEI (Proyecto Educativo)", "Manual de Convivencia", "Plan de Estudios"]
        }

    def _create_styled_docx(self, title, content, output_path):
        doc = Document()
        
        # Header Pro
        section = doc.sections[0]
        header = section.header
        htable = header.add_table(1, 2, width=Inches(6))
        htable.cell(0, 0).text = "HMO AUDITOR - SIMULACION ELITE"
        htable.cell(0, 1).text = f"REV: 01 | {datetime.date.today()}"
        
        # Título
        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contenido
        for subtitle, text in content.items():
            sh = doc.add_heading(subtitle, level=1)
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.save(output_path)
        return output_path

    def simulate_norm_ecosystem(self, norm_key, company_name, industry):
        """Pobla el expediente con documentos realistas para una norma específica."""
        if norm_key not in self.norms_config:
            return False, f"Norma {norm_key} no soportada."
            
        target_norm = self.norms_config[norm_key][0]
        docs_to_gen = self.norms_config[norm_key][1:]
        
        results = []
        for doc_name in docs_to_gen:
            # Usamos la IA para generar contenido REALISTA, no lorem ipsum
            print(f"Buscando IA para: {doc_name}...")
            raw_text = self.ai_engine.generate_draft_text(
                doc_name, company_name, industry, "Simulación de madurez corporativa HMO"
            )
            
            # Crear estructura de secciones básica para el Word
            content_map = {
                "1. OBJETIVO Y ALCANCE": f"Este documento establece los lineamientos para {doc_name} bajo la norma {target_norm}.",
                "2. DESARROLLO SUSTANCIAL": raw_text,
                "3. CONTROL Y SEGUIMIENTO": "Este activo digital es monitoreado por el Motor de Cumplimiento HMO Auditor."
            }
            
            # Guardar en la carpeta de la empresa
            folder = "01_Simulacion_Elite"
            os.makedirs(os.path.join(self.base_path, folder), exist_ok=True)
            filename = f"SIM_{doc_name.replace(' ', '_')}.docx".replace('(', '').replace(')', '')
            path = os.path.join(self.base_path, folder, filename)
            
            self._create_styled_docx(doc_name, content_map, path)
            results.append({"doc": doc_name, "path": path})
            
        return True, results

if __name__ == "__main__":
    # Prueba rápida
    engine = HMOSimulationEngine(r"d:\HMO\SENA\Auditor_Formatos\Simulacion_Demo")
    success, log = engine.simulate_norm_ecosystem("SEGURIDAD", "CiberGuard SAS", "Ciberseguridad")
    print(f"Simulación exitosa: {success}")
    if success:
        for r in log: print(f" - {r['doc']} -> {r['path']}")
